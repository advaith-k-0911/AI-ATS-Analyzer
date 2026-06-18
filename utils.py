import os
import json
import docx
from pypdf import PdfReader
from groq import Groq
from dotenv import load_dotenv
# Load environment variables
load_dotenv()

DEFAULT_MODEL = "llama-3.3-70b-versatile"
DEFAULT_TEMPERATURE = 0.2


def build_application_context(target_role=None, experience_level=None, focus_area=None):
    """Create a short context block from user-facing personalization choices."""
    context_lines = []

    if target_role and target_role.strip():
        context_lines.append(f"Target role: {target_role.strip()}")
    if experience_level and experience_level != "Not specified":
        context_lines.append(f"Experience level: {experience_level}")
    if focus_area:
        context_lines.append(f"Main focus requested by user: {focus_area}")

    return "\n".join(context_lines) if context_lines else "No extra user preferences provided."


def extract_text_from_pdf(file_path):
    """Extract text from a PDF file using pypdf."""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise Exception(f"Error reading PDF file: {str(e)}")
    return text
def extract_text_from_docx(file_path):
    """Extract text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")
    return text
def parse_resume_file(uploaded_file):
    """Parse the uploaded file based on its extension."""
    filename = uploaded_file.name
    # Save file temporarily or read bytes directly
    # Since streamlit upload is a BytesIO-like object, we can write a temporary file or pass it to standard tools
    # Let's save it temporarily to read it
    temp_dir = os.path.join(os.path.dirname(__file__), "temp")
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, filename)
    
    try:
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(temp_path)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(temp_path)
        else:
            # Fallback to plain text
            with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    finally:
        # Cleanup
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
    return text
def analyze_resume_ats(
    resume_text,
    job_desc_text,
    api_key=None,
    target_role=None,
    experience_level=None,
    focus_area=None
):
    """
    Review a resume against a job description and return a structured
    JSON payload for the Streamlit results screen.
    """
    # Use provided key or fall back to environment variable
    key = api_key or os.getenv("GROQ_API_KEY")
    if not key:
        raise ValueError("The analysis service is not configured yet.")
    client = Groq(api_key=key)
    application_context = build_application_context(target_role, experience_level, focus_area)
    system_prompt = (
        "You are an expert Applicant Tracking System (ATS) reviewer, HR consultant, and resume coach. "
        "Your task is to perform a highly thorough, accurate, and professional ATS audit of the user's resume against "
        "the provided job description. You must output the analysis strictly in JSON format. Do not write any conversational "
        "text before or after the JSON payload. Ensure all JSON fields are populated accurately. "
        "Keep the wording practical and applicant-friendly. Do not mention model names, API providers, temperature, prompts, "
        "JSON schema details, or backend implementation in any user-facing response field."
    )
    user_prompt = f"""
Analyze the resume and job description below.
Use the applicant preferences to make the recommendations more specific.
--- APPLICANT PREFERENCES ---
{application_context}
--- RESUME ---
{resume_text}
--- JOB DESCRIPTION ---
{job_desc_text}
--- OUTPUT JSON SCHEMA ---
You must return a JSON object with the following exact keys and types:
{{
  "ats_score": <int, 0 to 100 representing general compatibility>,
  "match_percentage": <int, 0 to 100 representing job description keyword/requirement alignment>,
  "formatting_score": <int, 0 to 100 representing layout/font/readability elements>,
  "keyword_score": <int, 0 to 100 representing presence of core resume keywords>,
  "experience_score": <int, 0 to 100 representing relevance of past roles to requirements>,
  "skills_score": <int, 0 to 100 representing job-related and soft skills alignment>,
  "analysis_summary": "<string, professional summary of candidate fit, highlighting core alignments and main gaps>",
  "key_strengths": ["<string>", "<string>", ...],
  "critical_issues": ["<string>", "<string>", ...],
  "keyword_analysis": [
    {{
      "keyword": "<string, keyword/skill name>",
      "status": "<string, either 'Present' or 'Missing'>",
      "importance": "<string, either 'High', 'Medium', or 'Low'>"
    }},
    ...
  ],
  "skills_gap": [
    {{
      "skill": "<string, skill/tool name>",
      "category": "<string, e.g. 'Hard Skill', 'Soft Skill', 'Tool/Technology'>",
      "gap_description": "<string, detail of how the requirement is missing or underrepresented in the resume>"
    }},
    ...
  ],
  "ats_compatibility_checks": [
    {{
      "check": "<string, name of standard check, e.g., 'Contact Info Found', 'Standard Headings Used', 'Bullet Points Formatting', 'No Tables/Columns'>",
      "status": "<string, 'Passed', 'Warning', or 'Failed'>",
      "feedback": "<string, constructive feedback on how to fix or why it passed/warned>"
    }},
    ...
  ],
  "improvement_suggestions": ["<string, clear actionable advice>", ...],
  "optimized_resume_snippets": [
    {{
      "original": "<string, sub-optimal bullet point or description in the resume>",
      "optimized": "<string, rewritten action-oriented bullet point that includes keywords and quantifies achievements>",
      "rationale": "<string, why the rewrite is better and what ATS keywords it addresses>"
    }},
    ...
  ],
  "interview_questions": [
    {{
      "question": "<string, tailored behavioral or role-specific question to ask this candidate during interview based on gaps>",
      "rationale": "<string, why this question is relevant to test the gap>"
    }},
    ...
  ],
  "cover_letter": "<string, a professional, compelling, and fully-formed cover letter matching the job description and leveraging candidate strengths from the resume. Use standard cover letter layout and placeholders for contact fields. Keep formatting clean with newlines.>"
}}
"""
    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=DEFAULT_TEMPERATURE
        )
        
        result_content = response.choices[0].message.content
        return json.loads(result_content)
    except json.JSONDecodeError:
        raise Exception("The analysis returned an unexpected format. Please try again.")
    except Exception:
        raise Exception("Could not complete the analysis right now. Please try again in a minute.")
